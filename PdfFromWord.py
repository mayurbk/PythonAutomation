## works only on windows and must have word installde in your machine.
import win32com.client
import docx
#wordfile = 'WordtoPDf.docx'
#pdffile = 'Mypdf.pdf'

writedoc = docx.Document()
writedoc.add_paragraph("this is our testing document.")

paraObj1 = writedoc.add_paragraph("this is our first paragraph.")
print(writedoc.save('WordtoPDf.docx'))
wdFormatPDF = 17 ## numeric code of word document for pdf
wordobj = win32com.client.Dispatch('Word.Application')
docObj = wordobj.Documents.Open('WordtoPDf.docx')
docObj.SaveAs('Mypdf.pdf', FileFormat = wdFormatPDF)
docObj.Close()
wordobj.Quit()