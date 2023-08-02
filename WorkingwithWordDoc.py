#working with word doc
import docx

import getDocxText

doc = docx.Document('demo.docx')
print(len(doc.paragraphs))

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


print(getText('demo.docx'))

#write documents
writedoc = docx.Document()
writedoc.add_paragraph("this is our testing document.")
#print(writedoc.save('Myowndocument.docx'))
paraObj1 = writedoc.add_paragraph("this is our first paragraph.")


#adding pictures to word document
writedoc.add_picture('zophie.png',width=docx.shared.Inches(1),height=docx.shared.Cm(4))
print(writedoc.save('Myowndocument.docx'))